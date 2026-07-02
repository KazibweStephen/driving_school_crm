"""
Fance Driving School - Integrated Driver Training Curriculum Document Generator
Generates a professional .docx with all 11 sections + appendix
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement
import datetime

doc = Document()

# ─── Colour Palette ──────────────────────────────────────────────
TEAL = RGBColor(0x00, 0x80, 0x80)
GOLD = RGBColor(0xD4, 0xAF, 0x37)
DARK_GOLD = RGBColor(0xB8, 0x96, 0x1E)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GREY_LIGHT = RGBColor(0xF2, 0xF2, 0xF2)
GREY_MED = RGBColor(0xE0, 0xE0, 0xE0)
GREY_DARK = RGBColor(0x99, 0x99, 0x99)
BLACK = RGBColor(0x33, 0x33, 0x33)
RED_SAFETY = RGBColor(0xCC, 0x33, 0x33)

# ─── Font Setup ──────────────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = BLACK
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for s_name in ['Heading 1', 'Heading 2', 'Heading 3']:
    s = doc.styles[s_name]
    s.font.name = 'Calibri'
    if s_name == 'Heading 1':
        s.font.size = Pt(18)
        s.font.bold = True
        s.font.color.rgb = TEAL
        s.paragraph_format.space_before = Pt(24)
        s.paragraph_format.space_after = Pt(12)
    elif s_name == 'Heading 2':
        s.font.size = Pt(14)
        s.font.bold = True
        s.font.color.rgb = GOLD
        s.paragraph_format.space_before = Pt(18)
        s.paragraph_format.space_after = Pt(8)
    elif s_name == 'Heading 3':
        s.font.size = Pt(12)
        s.font.bold = True
        s.font.color.rgb = TEAL
        s.paragraph_format.space_before = Pt(12)
        s.paragraph_format.space_after = Pt(6)

# ─── Helper Functions ────────────────────────────────────────────

def set_cell_shading(cell, color):
    """Set cell background color."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_text(cell, text, bold=False, color=BLACK, size=Pt(10), alignment=WD_ALIGN_PARAGRAPH.LEFT):
    """Set cell text with formatting."""
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = alignment
    run = p.add_run(text)
    run.bold = bold
    run.font.color.rgb = color
    run.font.size = size
    run.font.name = 'Calibri'
    # Reduce cell margins
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for margin in ['top', 'left', 'bottom', 'right']:
        elem = OxmlElement(f'w:{margin}')
        elem.set(qn('w:w'), '40')
        elem.set(qn('w:type'), 'dxa')
        tcMar.append(elem)
    tcPr.append(tcMar)

def set_cell_vertical_alignment(cell, align="center"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), align)
    tcPr.append(vAlign)

def add_styled_table(doc, headers, rows, col_widths=None, header_color='008080', alt_color='F2F2F2'):
    """Create a professionally styled table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, header_color)
        set_cell_text(cell, h, bold=True, color=WHITE, size=Pt(10), alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_vertical_alignment(cell, "center")

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            bg = alt_color if r_idx % 2 == 0 else 'FFFFFF'
            set_cell_shading(cell, bg)
            set_cell_text(cell, val, size=Pt(10))
            set_cell_vertical_alignment(cell, "center")

    # Set column widths if provided
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)

    return table

def add_section_break(doc, orientation=WD_ORIENT.PORTRAIT):
    """Add a section break (new page) with optional orientation."""
    new_section = doc.add_section()
    new_section.orientation = orientation
    if orientation == WD_ORIENT.LANDSCAPE:
        new_section.page_width = Cm(29.7)
        new_section.page_height = Cm(21)
    else:
        new_section.page_width = Cm(21)
        new_section.page_height = Cm(29.7)
    new_section.top_margin = Cm(2)
    new_section.bottom_margin = Cm(2)
    new_section.left_margin = Cm(2.5)
    new_section.right_margin = Cm(2.5)
    return new_section

def add_checkbox(doc, label, checked=False):
    """Add a checkbox line."""
    p = doc.add_paragraph()
    run = p.add_run('\u2611 ' if checked else '\u2610 ')
    run.font.size = Pt(11)
    run.font.name = 'Segoe UI Symbol'
    run.font.color.rgb = TEAL
    run2 = p.add_run(label)
    run2.font.size = Pt(11)
    run2.font.name = 'Calibri'
    return p

def add_instructor_tip(doc, text):
    """Call-out box for instructor tips (Gold accent)."""
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, 'FFF8E1')
    # Gold left border via XML
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for edge in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{edge}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '12')
        border.set(qn('w:color'), 'D4AF37')
        border.set(qn('w:space'), '0')
        borders.append(border)
    tcPr.append(borders)
    p = cell.paragraphs[0]
    run = p.add_run('\u2139\uFE0F  INSTRUCTOR TIP: ')
    run.bold = True
    run.font.color.rgb = DARK_GOLD
    run.font.size = Pt(10)
    run.font.name = 'Calibri'
    run2 = p.add_run(text)
    run2.font.size = Pt(10)
    run2.font.name = 'Calibri'
    doc.add_paragraph()  # spacer

def add_safety_note(doc, text):
    """Call-out box for safety information (Red accent)."""
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, 'FFF0F0')
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for edge in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{edge}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '12')
        border.set(qn('w:color'), 'CC3333')
        border.set(qn('w:space'), '0')
        borders.append(border)
    tcPr.append(borders)
    p = cell.paragraphs[0]
    run = p.add_run('\u26A0\uFE0F  IMPORTANT SAFETY: ')
    run.bold = True
    run.font.color.rgb = RED_SAFETY
    run.font.size = Pt(10)
    run.font.name = 'Calibri'
    run2 = p.add_run(text)
    run2.font.size = Pt(10)
    run2.font.name = 'Calibri'
    doc.add_paragraph()  # spacer

def add_icon_text(doc, icon, text, bold_prefix=None):
    """Add a line with an icon prefix."""
    p = doc.add_paragraph()
    run = p.add_run(f'{icon} ')
    run.font.size = Pt(12)
    run.font.name = 'Segoe UI Symbol'
    if bold_prefix:
        run2 = p.add_run(f'{bold_prefix}: ')
        run2.bold = True
        run2.font.size = Pt(11)
        run2.font.name = 'Calibri'
    run3 = p.add_run(text)
    run3.font.size = Pt(11)
    run3.font.name = 'Calibri'
    return p

def add_toc_field(doc):
    """Insert a TOC field."""
    p = doc.add_paragraph()
    run = p.add_run()
    fld_char_begin = OxmlElement('w:fldChar')
    fld_char_begin.set(qn('w:fldCharType'), 'begin')
    run._r.append(fld_char_begin)

    run2 = p.add_run()
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' TOC \\o "1-3" \\h \\z '
    run2._r.append(instr)

    run3 = p.add_run()
    fld_char_separate = OxmlElement('w:fldChar')
    fld_char_separate.set(qn('w:fldCharType'), 'separate')
    run3._r.append(fld_char_separate)

    run4 = p.add_run()
    fld_char_end = OxmlElement('w:fldChar')
    fld_char_end.set(qn('w:fldCharType'), 'end')
    run4._r.append(fld_char_end)

def add_toc_page(doc):
    """Add a dedicated TOC page."""
    p = doc.add_paragraph()
    run = p.add_run('TABLE OF CONTENTS')
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = TEAL
    run.font.name = 'Calibri'
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(20)

    # Gold line
    p2 = doc.add_paragraph()
    run2 = p2.add_run('_' * 70)
    run2.font.color.rgb = GOLD
    run2.font.size = Pt(8)
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_toc_field(doc)

    p3 = doc.add_paragraph()
    p3.paragraph_format.space_before = Pt(20)
    run3 = p3.add_run('(Right-click and select "Update Field" to populate the Table of Contents)')
    run3.font.size = Pt(9)
    run3.font.italic = True
    run3.font.color.rgb = GREY_DARK

def setup_header_footer(section, show=True):
    """Configure header and footer for a section."""
    if show:
        header = section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = hp.add_run('FANCE DRIVING SCHOOL \u2014 Integrated Driver Training Curriculum')
        run.font.size = Pt(8)
        run.font.color.rgb = TEAL
        run.font.name = 'Calibri'
        run.italic = True

        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Page number
        run_pg = fp.add_run('Page ')
        run_pg.font.size = Pt(8)
        run_pg.font.color.rgb = GREY_DARK
        run_pg.font.name = 'Calibri'

        # Insert PAGE field
        fld_char_begin = OxmlElement('w:fldChar')
        fld_char_begin.set(qn('w:fldCharType'), 'begin')
        run_fld = fp.add_run()
        run_fld._r.append(fld_char_begin)

        instr = OxmlElement('w:instrText')
        instr.set(qn('xml:space'), 'preserve')
        instr.text = ' PAGE '
        run_instr = fp.add_run()
        run_instr._r.append(instr)

        fld_char_sep = OxmlElement('w:fldChar')
        fld_char_sep.set(qn('w:fldCharType'), 'separate')
        run_sep = fp.add_run()
        run_sep._r.append(fld_char_sep)

        run0 = fp.add_run('1')
        run0.font.size = Pt(8)
        run0.font.color.rgb = GREY_DARK

        fld_char_end = OxmlElement('w:fldChar')
        fld_char_end.set(qn('w:fldCharType'), 'end')
        run_end = fp.add_run()
        run_end._r.append(fld_char_end)

        run_ver = fp.add_run('  |  Version 1.0')
        run_ver.font.size = Pt(8)
        run_ver.font.color.rgb = GREY_DARK
        run_ver.font.name = 'Calibri'
    else:
        header = section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0]
        hp.text = ''
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0]
        fp.text = ''

# ═══════════════════════════════════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════════════════════════════════

# Add several blank paragraphs before the title
for _ in range(6):
    doc.add_paragraph()

# Logo placeholder
p_logo = doc.add_paragraph()
p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_logo = p_logo.add_run('[ FANCE DRIVING SCHOOL LOGO ]')
run_logo.font.size = Pt(12)
run_logo.font.color.rgb = GREY_DARK
run_logo.font.italic = True
run_logo.font.name = 'Calibri'

doc.add_paragraph()

# Main title
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_t = p_title.add_run('FANCE DRIVING SCHOOL')
run_t.bold = True
run_t.font.size = Pt(32)
run_t.font.color.rgb = TEAL
run_t.font.name = 'Calibri'

doc.add_paragraph()

# Subtitle
p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_sub = p_sub.add_run('Integrated Driver Training Curriculum')
run_sub.font.size = Pt(20)
run_sub.font.color.rgb = GOLD
run_sub.font.name = 'Calibri'

p_sub2 = doc.add_paragraph()
p_sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_sub2 = p_sub2.add_run('Manual & Automatic Transmission')
run_sub2.font.size = Pt(16)
run_sub2.font.color.rgb = TEAL
run_sub2.font.name = 'Calibri'

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

# Gold divider
p_div = doc.add_paragraph()
p_div.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_div = p_div.add_run('\u2501' * 50)
run_div.font.color.rgb = GOLD
run_div.font.size = Pt(10)

doc.add_paragraph()

# Document type
p_type = doc.add_paragraph()
p_type.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_type = p_type.add_run('Client Proposal  |  Instructor Training Manual  |  Curriculum Guide  |  CRM Functional Specification')
run_type.font.size = Pt(11)
run_type.font.color.rgb = BLACK
run_type.font.name = 'Calibri'

doc.add_paragraph()
doc.add_paragraph()

# Version and date
p_ver = doc.add_paragraph()
p_ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_ver = p_ver.add_run(f'Version 1.0  |  {datetime.date.today().strftime("%B %Y")}')
run_ver.font.size = Pt(12)
run_ver.font.color.rgb = GREY_DARK
run_ver.font.name = 'Calibri'

p_prep = doc.add_paragraph()
p_prep.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_prep = p_prep.add_run('Prepared for Fance Driving School Management')
run_prep.font.size = Pt(11)
run_prep.font.color.rgb = GREY_DARK
run_prep.font.name = 'Calibri'

# Add section break after cover
sec_cover = doc.sections[0]
setup_header_footer(sec_cover, show=False)

# ═══════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ═══════════════════════════════════════════════════════════════════

new_section = add_section_break(doc, WD_ORIENT.PORTRAIT)
setup_header_footer(new_section, show=True)

p_toc_title = doc.add_paragraph()
run_toc_t = p_toc_title.add_run('TABLE OF CONTENTS')
run_toc_t.bold = True
run_toc_t.font.size = Pt(20)
run_toc_t.font.color.rgb = TEAL
run_toc_t.font.name = 'Calibri'
p_toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_toc_title.paragraph_format.space_after = Pt(10)

p_toc_line = doc.add_paragraph()
run_toc_line = p_toc_line.add_run('_' * 70)
run_toc_line.font.color.rgb = GOLD
run_toc_line.font.size = Pt(8)
p_toc_line.alignment = WD_ALIGN_PARAGRAPH.CENTER

add_toc_field(doc)

p_toc_note = doc.add_paragraph()
p_toc_note.paragraph_format.space_before = Pt(20)
run_toc_note = p_toc_note.add_run('(Right-click and select "Update Field" to populate the Table of Contents)')
run_toc_note.font.size = Pt(9)
run_toc_note.font.italic = True
run_toc_note.font.color.rgb = GREY_DARK

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# SECTION 1: COURSE OVERVIEW
# ═══════════════════════════════════════════════════════════════════

doc.add_heading('1. Course Overview', level=1)

p_intro = doc.add_paragraph()
run_intro = p_intro.add_run(
    'The Fance Driving School Integrated Driver Training Programme is a comprehensive, '
    'competency-based driver education course designed for both Manual and Automatic transmission vehicles. '
    'This programme combines practical behind-the-wheel training with theoretical classroom instruction '
    'to produce safe, confident, and responsible drivers.'
)
run_intro.font.size = Pt(11)
run_intro.font.name = 'Calibri'

doc.add_paragraph()

overview_headers = ['Item', 'Details']
overview_rows = [
    ['Course Name', 'Integrated Driver Training Programme'],
    ['Vehicle Types', 'Manual & Automatic'],
    ['Practical Lessons', '20 Lessons (4 Weeks \u00d7 5 Days)'],
    ['Theory Lessons', 'Every Saturday (4 Sessions \u00d7 2 Hours)'],
    ['Practical Duration', '30 Minutes OR 3 km OR Until Lesson Competencies Are Achieved (Whichever Comes First)'],
    ['Theory Duration', '2 Hours per Session'],
    ['Assessment', 'Competency-Based Assessment'],
    ['Class Size', 'One-on-One (Instructor : Student)'],
    ['Target Audience', 'Beginner Drivers (No Prior Experience Required)'],
    ['Certification', 'Certificate of Completion \u2013 Fance Driving School'],
]

add_styled_table(doc, overview_headers, overview_rows, col_widths=[6, 12])

doc.add_paragraph()

add_icon_text(doc, '\u2714\ufe0f', ' All practical lessons are one-on-one with a qualified instructor.')
add_icon_text(doc, '\u2714\ufe0f', ' Theory sessions are conducted every Saturday in a classroom setting.')
add_icon_text(doc, '\u2714\ufe0f', ' Both Manual and Automatic transmission options are available.')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# SECTION 2: TRAINING PHILOSOPHY
# ═══════════════════════════════════════════════════════════════════

doc.add_heading('2. Training Philosophy', level=1)

doc.add_heading('Competency-Based Learning', level=2)

p_phil = doc.add_paragraph()
run_phil = p_phil.add_run(
    'The Fance Driving School curriculum is built on a competency-based learning model. '
    'Unlike traditional time-based driver training where a fixed number of hours is mandated, '
    'our approach focuses on the demonstration of specific, measurable competencies before '
    'a student progresses to the next lesson. This ensures that every student masters the '
    'essential skills required for safe and responsible driving.'
)
run_phil.font.size = Pt(11)
run_phil.font.name = 'Calibri'

doc.add_paragraph()
doc.add_heading('Key Principles', level=2)

principles = [
    ('\u2714\ufe0f  Mastery Before Progression', 'Students advance to the next lesson only after demonstrating competency in all required skills for the current lesson.'),
    ('\u2714\ufe0f  Flexible Lesson Duration', 'Lessons may conclude before the 30-minute mark if the student has achieved all lesson competencies. Conversely, lessons may be extended or repeated if competencies are not yet met.'),
    ('\u2714\ufe0f  Individualised Pacing', 'Each student learns at their own pace. There is no fixed schedule for completing the 20-lesson programme.'),
    ('\u2714\ufe0f  Continuous Assessment', 'Every lesson includes a formal assessment against a predefined checklist of competencies.'),
    ('\u2714\ufe0f  Instructor Discretion', 'Instructors have the professional discretion to repeat or modify lessons based on student progress and road conditions.'),
    ('\u2714\ufe0f  Integrated Feedback', 'Students receive immediate, constructive feedback after every lesson to reinforce learning and address weaknesses.'),
]

for title, desc in principles:
    p = doc.add_paragraph()
    run_t = p.add_run(f'{title}\n')
    run_t.bold = True
    run_t.font.size = Pt(11)
    run_t.font.name = 'Calibri'
    run_d = p.add_run(desc)
    run_d.font.size = Pt(11)
    run_d.font.name = 'Calibri'
    p.paragraph_format.space_after = Pt(6)

doc.add_paragraph()
doc.add_heading('Competency Definition', level=2)

p_comp = doc.add_paragraph()
run_comp = p_comp.add_run(
    'A competency is defined as a specific, observable, and measurable skill that a student must '
    'demonstrate consistently and without assistance. Competencies are assessed on a pass/fail basis '
    'using a standardised instructor checklist. All competencies must be assessed as "Achieved" for '
    'the lesson to be marked complete.'
)
run_comp.font.size = Pt(11)
run_comp.font.name = 'Calibri'

add_instructor_tip(doc, 'Always provide verbal feedback during the assessment. Competency-based learning is not about pass/fail alone \u2014 it is about continuous improvement and building confident drivers.')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# SECTION 3: DAILY PRACTICAL SESSION STRUCTURE
# ═══════════════════════════════════════════════════════════════════

doc.add_heading('3. Daily Practical Session Structure', level=1)

p_session = doc.add_paragraph()
run_session = p_session.add_run(
    'Every practical driving lesson follows a standardised 30-minute session structure. '
    'This ensures consistency across instructors and provides students with a predictable '
    'learning routine. Each session comprises five distinct phases:'
)
run_session.font.size = Pt(11)
run_session.font.name = 'Calibri'

doc.add_paragraph()

session_headers = ['Activity', 'Duration', 'Description']
session_rows = [
    ['Vehicle Safety Inspection', '3 Minutes', 'Daily vehicle inspection conducted by student under instructor supervision. Covers tyres, lights, fluids, and general vehicle safety.'],
    ['Cockpit Drill (DSSSM)', '3 Minutes', 'Doors, Seat, Steering, Seatbelt, Mirrors. Student demonstrates correct driving position setup.'],
    ['Lesson Video', '2\u20133 Minutes', 'Short instructional video relevant to the day\'s lesson objectives. Viewed on tablet or in-vehicle display.'],
    ['Practical Session', '20 Minutes', 'Hands-on driving practice focusing on the lesson\'s specific competencies. Instructor provides real-time guidance and feedback.'],
    ['Daily Assessment', '4 Minutes', 'Instructor completes the daily checklist, reviews competencies achieved, provides verbal feedback, and records progress in the CRM.'],
]

add_styled_table(doc, session_headers, session_rows, col_widths=[4, 3, 11])

doc.add_paragraph()

# Highlight box for session duration
p_highlight = doc.add_paragraph()
p_highlight.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_highlight.paragraph_format.space_before = Pt(12)
p_highlight.paragraph_format.space_after = Pt(6)
run_hl = p_highlight.add_run('SESSION DURATION')
run_hl.bold = True
run_hl.font.size = Pt(14)
run_hl.font.color.rgb = TEAL
run_hl.font.name = 'Calibri'

# Create a highlight table
labels = [
    ['30 Minutes', 'OR', '3 Kilometres'],
    ['', 'OR', ''],
    ['Lesson Competencies Achieved', '', ''],
    ['Whichever Comes First', '', ''],
]
hl_table = doc.add_table(rows=4, cols=3)
hl_table.alignment = WD_TABLE_ALIGNMENT.CENTER
hl_table.style = 'Table Grid'

hl_table.rows[0].cells[0].merge(hl_table.rows[0].cells[2])
hl_table.rows[1].cells[0].merge(hl_table.rows[1].cells[2])
hl_table.rows[2].cells[0].merge(hl_table.rows[2].cells[2])
hl_table.rows[3].cells[0].merge(hl_table.rows[3].cells[2])

for r in range(4):
    for c in range(3):
        cell = hl_table.rows[r].cells[c]
        set_cell_shading(cell, 'FFF8E1')
        val = labels[r][c]
        is_or = (val == 'OR')
        is_big = val in ['30 Minutes', '3 Kilometres', 'Lesson Competencies Achieved']
        set_cell_text(cell, val,
                     bold=is_or or val == 'Whichever Comes First',
                     color=TEAL if is_big else GOLD,
                     size=Pt(16) if is_big else Pt(12),
                     alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_vertical_alignment(cell, "center")

doc.add_paragraph()

add_safety_note(doc, 'The "Whichever Comes First" rule ensures that no student exceeds 30 minutes of continuous driving or 3 km without a formal assessment break. This prevents fatigue and maintains high safety standards.')

add_instructor_tip(doc, 'For students who complete competencies early, use the remaining session time for revision of previous lessons or introduction of the next lesson\'s concepts.')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# SECTION 4: TWENTY-DAY PRACTICAL CURRICULUM
# ═══════════════════════════════════════════════════════════════════

# Switch to landscape for the curriculum table
landscape_section = add_section_break(doc, WD_ORIENT.LANDSCAPE)
setup_header_footer(landscape_section, show=True)

doc.add_heading('4. Twenty-Day Practical Curriculum', level=1)

p_curr_intro = doc.add_paragraph()
run_curr = p_curr_intro.add_run(
    'The practical curriculum is organised across four weeks (five lessons per week). '
    'Each lesson has defined competencies, learning objectives, and assessment criteria. '
    'The table below presents the complete curriculum for both Manual and Automatic transmission vehicles.'
)
run_curr.font.size = Pt(11)
run_curr.font.name = 'Calibri'

doc.add_paragraph()

# Define the 20-day curriculum data
curriculum_data = [
    {
        'day': 1, 'week': 1,
        'title': 'Vehicle Orientation',
        'manual': 'Yes', 'auto': 'Yes',
        'competencies': 'Identify all vehicle controls; Adjust seat, steering, mirrors; Locate gear positions (Manual); Identify PRNDL (Auto)',
        'video': 'Vehicle Orientation & Controls',
        'objectives': 'Familiarise student with vehicle interior, controls, dashboard indicators, and basic adjustments.',
        'checklist': [
            'Identifies steering wheel, pedals, gear lever, handbrake',
            'Adjusts driver seat correctly',
            'Adjusts steering wheel position',
            'Adjusts rear-view and side mirrors',
            'Identifies dashboard warning lights',
            'Demonstrates understanding of gear positions',
            'Shows awareness of vehicle dimensions',
        ],
        'assessment': 'Student can identify and demonstrate all primary controls without assistance.',
        'crm_status': 'Day 1 Orientation Complete',
        'notes': 'This is a non-driving lesson. Focus on building familiarity and confidence. Do not start the engine for driving purposes.',
    },
    {
        'day': 2, 'week': 1,
        'title': 'Vehicle Checks',
        'manual': 'Yes', 'auto': 'Yes',
        'competencies': 'Perform daily vehicle inspection; Check tyre condition; Verify fluid levels; Demonstrate safety awareness',
        'video': 'Daily Vehicle Inspection',
        'objectives': 'Teach the student how to perform a thorough daily vehicle safety inspection.',
        'checklist': [
            'Performs walk-around inspection',
            'Checks tyre condition and pressure (visual)',
            'Verifies headlights, tail lights, indicators',
            'Checks brake lights',
            'Verifies engine oil level',
            'Checks coolant level',
            'Checks windshield washer fluid',
            'Demonstrates safety awareness throughout',
        ],
        'assessment': 'Student completes full vehicle inspection checklist independently.',
        'crm_status': 'Day 2 Vehicle Checks Complete',
        'notes': 'Use the Vehicle Inspection Checklist in the Appendix. Emphasise safety-first approach.',
    },
    {
        'day': 3, 'week': 1,
        'title': 'Clutch Control (Manual) / Automatic Controls',
        'manual': 'Yes', 'auto': 'Yes',
        'competencies': 'Find bite point (Manual); Smooth acceleration (Auto); Controlled braking; Steering while stationary',
        'video': 'Clutch Control Basics / Automatic Controls',
        'objectives': 'Develop clutch control (Manual) or accelerator/brake modulation (Auto) in a controlled environment.',
        'checklist': [
            '[Manual] Finds bite point consistently',
            '[Manual] Moves off without stalling',
            '[Manual] Changes gear smoothly (1st \u2192 2nd)',
            '[Auto] Selects D, R, P correctly',
            '[Auto] Accelerates smoothly',
            '[Auto] Brakes smoothly to stop',
            'Steers while moving at low speed',
            'Uses mirrors before moving off',
        ],
        'assessment': 'Student demonstrates smooth control of vehicle at low speed in a safe, traffic-free area.',
        'crm_status': 'Day 3 Controls Practised',
        'notes': 'Conduct this lesson in an empty parking lot or quiet street. Prioritise smoothness over speed.',
    },
    {
        'day': 4, 'week': 1,
        'title': 'Steering',
        'manual': 'Yes', 'auto': 'Yes',
        'competencies': 'Correct steering technique; Pull-push method; Steering accuracy; Low-speed manoeuvring',
        'video': 'Steering Techniques',
        'objectives': 'Develop proper steering technique using the pull-push method. Maintain vehicle position during turns.',
        'checklist': [
            'Demonstrates pull-push steering technique',
            'Maintains correct hand position (10 & 2 or 9 & 3)',
            'Steers smoothly through gentle turns',
            'Avoids crossing hands',
            'Returns steering wheel correctly after turns',
            'Maintains lane position while steering',
            'Steers at low speed with precision',
        ],
        'assessment': 'Student executes turns using correct pull-push technique without crossing hands.',
        'crm_status': 'Day 4 Steering Practised',
        'notes': 'Emphasise pull-push method exclusively. Correct any hand-crossing habits immediately.',
    },
    {
        'day': 5, 'week': 1,
        'title': 'Consolidation',
        'manual': 'Yes', 'auto': 'Yes',
        'competencies': 'All week 1 competencies; Smooth driving at low speed; Basic observation',
        'video': 'Week 1 Consolidation',
        'objectives': 'Reinforce and consolidate all skills learned in Week 1. Identify and address any weaknesses.',
        'checklist': [
            'Demonstrates vehicle controls confidently',
            'Performs vehicle inspection independently',
            'Applies correct cockpit drill',
            'Controls vehicle smoothly at low speed',
            'Steers correctly using pull-push method',
            '[Manual] Changes gear without looking',
            '[Auto] Demonstrates smooth pedal control',
            'Shows basic observation awareness',
        ],
        'assessment': 'Satisfactory demonstration of all Week 1 competencies. Any gaps trigger a repeat lesson.',
        'crm_status': 'Week 1 Consolidated',
        'notes': 'Use this lesson to assess overall progress. If the student struggles with any key skill, schedule a repeat before Week 2.',
    },
    {
        'day': 6, 'week': 2,
        'title': 'Moving Off',
        'manual': 'Yes', 'auto': 'Yes',
        'competencies': 'Safe moving off; Mirror checks; Signal use; Correct gear selection; Observation at junctions',
        'video': 'Moving Off & Stopping',
        'objectives': 'Teach the student to move off safely from a parked position, including all necessary observations and signals.',
        'checklist': [
            'Performs full observation before moving off',
            'Checks mirrors (interior, left, right)',
            'Signals intention correctly',
            'Selects appropriate gear',
            'Moves off smoothly and safely',
            'Checks blind spot',
            'Stops smoothly at designated point',
            'Applies handbrake when stopped',
        ],
        'assessment': 'Student moves off and stops safely and smoothly with correct observation sequence.',
        'crm_status': 'Day 6 Moving Off Practised',
        'notes': 'Practice on a quiet residential street. Emphasise the MSM (Mirror-Signal-Manoeuvre) routine.',
    },
    {
        'day': 7, 'week': 2,
        'title': 'Junctions',
        'manual': 'Yes', 'auto': 'Yes',
        'competencies': 'Approach junctions safely; Correct gear selection; Observation at junctions; Gap selection; Yielding right of way',
        'video': 'Approaching Junctions',
        'objectives': 'Develop the ability to approach, assess, and navigate both controlled and uncontrolled junctions safely.',
        'checklist': [
            'Identifies junction type (controlled/uncontrolled)',
            'Approaches at appropriate speed',
            'Selects correct gear for approach',
            'Performs effective observation',
            'Judges gaps correctly',
            'Yields right of way appropriately',
            'Emerges safely from junction',
            'Checks mirrors after emerging',
        ],
        'assessment': 'Student navigates junctions safely with appropriate speed, gear, and observation.',
        'crm_status': 'Day 7 Junctions Practised',
        'notes': 'Start with simple T-junctions before progressing to crossroads and complex junctions.',
    },
    {
        'day': 8, 'week': 2,
        'title': 'Hill Starts & Hill Parking',
        'manual': 'Yes', 'auto': 'Yes',
        'competencies': 'Hill start (Manual: handbrake method); Hill start (Auto); Parking on incline; Restart on incline',
        'video': 'Hill Starts & Incline Parking',
        'objectives': 'Teach the student to start, stop, and park safely on inclines for both Manual and Automatic vehicles.',
        'checklist': [
            '[Manual] Uses handbrake for hill start',
            '[Manual] Finds bite point on hill',
            '[Manual] Moves off without rolling back',
            '[Auto] Uses handbrake on steep hill',
            '[Auto] Accelerates smoothly on incline',
            'Parks facing uphill with wheels turned',
            'Parks facing downhill with wheels turned',
            'Restarts smoothly on incline',
        ],
        'assessment': 'Student performs hill start without rolling back and parks correctly on incline with wheels positioned properly.',
        'crm_status': 'Day 8 Hill Skills Practised',
        'notes': 'Choose a moderate incline first, then progress to steeper hills. Check local regulations for wheel positioning.',
    },
    {
        'day': 9, 'week': 2,
        'title': 'Reverse Driving & Reverse Parking',
        'manual': 'Yes', 'auto': 'Yes',
        'competencies': 'Reverse in a straight line; Steer while reversing; Reverse parking between two cars; Observation during reversing',
        'video': 'Reverse Driving & Parking',
        'objectives': 'Develop confidence and control in reverse, including parking between two vehicles.',
        'checklist': [
            'Performs rear observation before reversing',
            'Reverses in straight line',
            'Steers accurately while reversing',
            'Maintains slow, controlled speed',
            'Parks between two cars (reverse bay)',
            'Corrects position if needed',
            'Uses mirrors effectively while reversing',
            'Checks blind spots during manoeuvre',
        ],
        'assessment': 'Student reverses into a parking bay between two cars with minimal correction.',
        'crm_status': 'Day 9 Reverse Skills Practised',
        'notes': 'Use cones initially if the student lacks confidence. Build up to parking between actual vehicles.',
    },
    {
        'day': 10, 'week': 2,
        'title': 'Bay Parking & Angle Parking',
        'manual': 'Yes', 'auto': 'Yes',
        'competencies': 'Forward bay parking; Reverse bay parking; Angle parking; Observation; Accuracy',
        'video': 'Bay & Angle Parking',
        'objectives': 'Master both forward and reverse bay parking, as well as angle parking, with precision and confidence.',
        'checklist': [
            'Performs forward bay parking',
            'Performs reverse bay parking',
            'Parks at an angle correctly',
            'Maintains safe distance from adjacent cars',
            'Uses reference points consistently',
            'Corrects position as needed',
            'Demonstrates full observation throughout',
            'Completes parking within reasonable attempts',
        ],
        'assessment': 'Student completes forward bay, reverse bay, and angle parking with satisfactory accuracy.',
        'crm_status': 'Day 10 Parking Practised',
        'notes': 'Teach reference points specific to the training vehicle. These may differ between Manual and Automatic due to different dimensions.',
    },
    {
        'day': 11, 'week': 3,
        'title': 'Road Positioning',
        'manual': 'Yes', 'auto': 'Yes',
        'competencies': 'Correct lane position; Positioning for turns; Overtaking position; Defensive positioning',
        'video': 'Road Positioning & Lane Discipline',
        'objectives': 'Develop proper road positioning skills for straight driving, turns, and general lane discipline.',
        'checklist': [
            'Maintains centre lane position on straight roads',
            'Positions correctly for left turns',
            'Positions correctly for right turns',
            'Adjusts position for oncoming vehicles',
            'Maintains safe distance from parked cars',
            'Uses correct lane for destination',
            'Signals lane changes appropriately',
            'Checks mirrors before changing position',
        ],
        'assessment': 'Student maintains correct road position consistently across different road types and conditions.',
        'crm_status': 'Day 11 Positioning Practised',
        'notes': 'Drive on a variety of road types: residential, arterial, and multi-lane roads where available.',
    },
    {
        'day': 12, 'week': 3,
        'title': 'Roundabouts',
        'manual': 'Yes', 'auto': 'Yes',
        'competencies': 'Approach roundabouts; Lane selection; Yield rules; Exit positioning; Multi-lane roundabouts',
        'video': 'Navigating Roundabouts',
        'objectives': 'Teach the student to navigate single-lane and multi-lane roundabouts safely and confidently.',
        'checklist': [
            'Approaches roundabout at appropriate speed',
            'Selects correct lane for intended exit',
            'Yields to traffic from the right',
            'Signals entry correctly',
            'Signals exit correctly',
            'Maintains lane on multi-lane roundabout',
            'Exits safely without cutting off others',
            'Checks mirrors after exiting',
        ],
        'assessment': 'Student navigates roundabouts safely with correct lane selection, signalling, and yielding.',
        'crm_status': 'Day 12 Roundabouts Practised',
        'notes': 'Start with mini-roundabouts, progress to single-lane, then multi-lane. Practice both left and right exits.',
    },
    {
        'day': 13, 'week': 3,
        'title': 'Parallel Parking Practice (Training Pitch)',
        'manual': 'Yes', 'auto': 'Yes',
        'competencies': 'Parallel parking procedure; Reference points; Observation; Accuracy; Correction techniques',
        'video': 'Parallel Parking Step by Step',
        'objectives': 'Develop parallel parking skills using a marked training pitch with reference points.',
        'checklist': [
            'Selects appropriate parking position',
            'Signals intention to park',
            'Performs full observation before manoeuvre',
            'Aligns vehicle correctly alongside front car',
            'Steers full lock at correct point',
            'Checks rear clearance during manoeuvre',
            'Straightens wheels at correct point',
            'Completes parking within two car lengths',
        ],
        'assessment': 'Student completes parallel parking manoeuvre within two car lengths with no more than one correction.',
        'crm_status': 'Day 13 Parallel Parking Practised',
        'notes': 'Use clearly marked reference points on the training pitch. The same reference points are used for Day 18 assessment.',
    },
    {
        'day': 14, 'week': 3,
        'title': 'Turning Around Obstacles & Safe Overtaking',
        'manual': 'Yes', 'auto': 'Yes',
        'competencies': 'Three-point turn; U-turn; Overtaking cyclist/parked car; Observation; Safety judgement',
        'video': 'Turning Around & Overtaking',
        'objectives': 'Develop the ability to turn the vehicle around safely and overtake stationary obstacles with confidence.',
        'checklist': [
            'Selects safe location for turn',
            'Performs three-point turn smoothly',
            'Performs U-turn where permitted',
            'Full observation before manoeuvre',
            'Overtakes stationary car safely',
            'Overtakes cyclist with sufficient clearance',
            'Checks blind spot before overtaking',
            'Returns to lane position smoothly',
        ],
        'assessment': 'Student completes three-point turn safely and overtakes obstacles with proper clearance and observation.',
        'crm_status': 'Day 14 Manoeuvres Practised',
        'notes': 'Emphasise the 1.5-metre clearance rule when overtaking cyclists. Check local road regulations.',
    },
    {
        'day': 15, 'week': 3,
        'title': 'Emergency Procedures',
        'manual': 'Yes', 'auto': 'Yes',
        'competencies': 'Emergency stop; Hazard avoidance; Brake failure procedure; Tyre burst awareness; Accident protocol',
        'video': 'Emergency Procedures',
        'objectives': 'Prepare students to respond correctly to emergency situations including emergency stops and vehicle failures.',
        'checklist': [
            'Performs emergency stop on command',
            'Maintains steering control during emergency stop',
            'Checks mirrors before braking',
            'Demonstrates hazard avoidance manoeuvre',
            'Explains brake failure procedure',
            'Explains tyre burst procedure',
            'Demonstrates accident protocol (stop, assess, call)',
            'Remains calm throughout',
        ],
        'assessment': 'Student performs controlled emergency stop and explains procedures for common emergencies.',
        'crm_status': 'Day 15 Emergency Procedures Complete',
        'notes': 'Ensure a safe, traffic-free area for emergency stop practice. Warm the student before performing the emergency stop.',
    },
    {
        'day': 16, 'week': 4,
        'title': 'Defensive Driving',
        'manual': 'Yes', 'auto': 'Yes',
        'competencies': 'Hazard perception; Following distance; Scanning; Anticipation; Speed management',
        'video': 'Defensive Driving Techniques',
        'objectives': 'Develop defensive driving habits: scanning, anticipation, maintaining safe following distance, and hazard perception.',
        'checklist': [
            'Maintains safe following distance (2-second rule)',
            'Scans road ahead continuously',
            'Anticipates potential hazards',
            'Adjusts speed for road conditions',
            'Positions vehicle defensively',
            'Identifies hidden hazards (e.g., driveways)',
            'Maintains awareness of surrounding traffic',
            'Avoids unnecessary risks',
        ],
        'assessment': 'Student demonstrates defensive driving habits consistently across varied road conditions.',
        'crm_status': 'Day 16 Defensive Driving Practised',
        'notes': 'This is a mindset lesson. Emphasise that defensive driving prevents accidents before they happen.',
    },
    {
        'day': 17, 'week': 4,
        'title': 'Mixed Road Driving',
        'manual': 'Yes', 'auto': 'Yes',
        'competencies': 'Highway driving; Urban driving; Rural driving; Speed adaptation; Lane merging',
        'video': 'Mixed Road Driving Strategies',
        'objectives': 'Give students experience driving on different road types: highways, urban streets, and rural roads.',
        'checklist': [
            'Merges onto highway safely',
            'Maintains appropriate highway speed',
            'Exits highway safely',
            'Navigates urban traffic confidently',
            'Adapts speed for rural road conditions',
            'Handles narrow roads with oncoming traffic',
            'Uses passing places correctly',
            'Demonstrates confident lane changes',
        ],
        'assessment': 'Student drives confidently and safely on highways, urban, and rural roads with appropriate speed adaptation.',
        'crm_status': 'Day 17 Mixed Road Driving Practised',
        'notes': 'Plan a route that includes all three road types. Adjust based on local road availability.',
    },
    {
        'day': 18, 'week': 4,
        'title': 'Parallel Parking Assessment (Training Pitch)',
        'manual': 'Yes', 'auto': 'Yes',
        'competencies': 'Parallel parking mastery; Timely execution; Accuracy; Observation; Independent performance',
        'video': 'Parallel Parking \u2013 Assessment Preparation',
        'objectives': 'Assess mastery of parallel parking using the same training pitch reference points from Day 13.',
        'checklist': [
            'Parks within two car lengths',
            'Completes manoeuvre in reasonable time',
            'No more than one correction',
            'Full observation throughout',
            'Does not mount kerb',
            'Parks parallel within 30 cm of kerb',
            'Performs independently without prompting',
            'Demonstrates confidence throughout',
        ],
        'assessment': 'Student passes parallel parking assessment against all criteria above. A fail requires additional practice.',
        'crm_status': 'Day 18 Parallel Parking Practised',
        'notes': 'This is an assessment-only lesson. Do not provide guidance unless safety is compromised.',
    },
    {
        'day': 19, 'week': 4,
        'title': 'Mock Practical Test',
        'manual': 'Yes', 'auto': 'Yes',
        'competencies': 'All competencies; Test procedures; Independent driving; Time management; Stress management',
        'video': 'Mock Test Preparation',
        'objectives': 'Simulate the official practical driving test to prepare the student for the real examination.',
        'checklist': [
            'Performs independent driving section',
            'Follows test route without prompting',
            'Demonstrates all manoeuvres on demand',
            'Manages time effectively',
            'Maintains composure under test conditions',
            'Recovers from minor errors',
            'Demonstrates test-ready competence',
            'Self-assesses performance accurately',
        ],
        'assessment': 'Student achieves mock test pass standard. Specific feedback provided on remaining weaknesses.',
        'crm_status': 'Day 19 Mock Test Complete',
        'notes': 'Use the official test route if known. Follow examiner procedures including timing and scoring.',
    },
    {
        'day': 20, 'week': 4,
        'title': 'Final Assessment',
        'manual': 'Yes', 'auto': 'Yes',
        'competencies': 'All 20-day competencies; Overall driving competence; Road safety awareness; Independent driving',
        'video': 'Final Assessment Review',
        'objectives': 'Conduct a comprehensive final assessment of all driving competencies. Determine readiness for licensing test.',
        'checklist': [
            'Demonstrates all competencies from Days 1\u201319',
            'Drives independently without prompting',
            'Applies defensive driving habits consistently',
            'Handles all road types confidently',
            'Parks using any method correctly',
            'Performs emergency stop on demand',
            'Demonstrates hazard perception throughout',
            'Achieves overall pass standard',
        ],
        'assessment': 'Student passes Final Assessment. Certificate of Completion is issued. Student is recommended for licensing test.',
        'crm_status': 'Day 20 Pass \u2013 Programme Complete',
        'notes': 'This is the final gate before the licensing test. Only recommend if the student is genuinely ready.',
    },
]

for lesson in curriculum_data:
    d = lesson['day']
    w = lesson['week']
    title = lesson['title']
    comps = lesson['competencies']
    video = lesson['video']
    obj = lesson['objectives']
    assessment = lesson['assessment']
    crm = lesson['crm_status']
    notes = lesson['notes']

    # Day header
    p_day = doc.add_paragraph()
    run_day = p_day.add_run(f'DAY {d}  |  Week {w}  |  {title}')
    run_day.bold = True
    run_day.font.size = Pt(12)
    run_day.font.color.rgb = TEAL
    run_day.font.name = 'Calibri'

    # Info table for this day
    day_headers = ['Field', 'Details']
    day_rows = [
        ['Competencies', comps],
        ['Learning Objectives', obj],
        ['Video Illustration', video],
        ['Assessment', assessment],
        ['CRM Status', crm],
        ['Instructor Notes', notes],
    ]
    add_styled_table(doc, day_headers, day_rows, col_widths=[4, 14])

    # Instructor checklist
    p_cl = doc.add_paragraph()
    run_cl = p_cl.add_run('Instructor Checklist:')
    run_cl.bold = True
    run_cl.font.size = Pt(10)
    run_cl.font.color.rgb = GOLD
    run_cl.font.name = 'Calibri'

    for item in lesson['checklist']:
        add_checkbox(doc, item)

    doc.add_paragraph()  # spacer

    # Add page break between weeks (except within same week)
    if d in [5, 10, 15]:
        doc.add_page_break()

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# SECTION 5: SATURDAY THEORY CURRICULUM
# ═══════════════════════════════════════════════════════════════════

# Switch back to portrait
portrait_section = add_section_break(doc, WD_ORIENT.PORTRAIT)
setup_header_footer(portrait_section, show=True)

doc.add_heading('5. Weekly Saturday Theory Sessions', level=1)

p_theory_intro = doc.add_paragraph()
run_ti = p_theory_intro.add_run(
    'Theory lessons are conducted every Saturday for four consecutive weeks. Each session is two hours long '
    'and covers essential theoretical knowledge required for safe driving. Students must attend all four sessions '
    'to qualify for the programme certificate.'
)
run_ti.font.size = Pt(11)
run_ti.font.name = 'Calibri'

doc.add_paragraph()

saturdays = [
    {
        'sat': 'Saturday 1 \nIntroduction & Fundamentals',
        'theme': 'Driver Responsibilities & Vehicle Familiarisation',
        'topics': [
            'Driver responsibilities and legal obligations',
            'Understanding different road users and their rights',
            'Cockpit drill (DSSSM) theory',
            'Dashboard warning lights interpretation',
            'Basic vehicle mechanics and maintenance awareness',
            'Pre-drive vehicle checks theory',
        ],
        'materials': ['Whiteboard/Projector', 'Vehicle cutaway diagram', 'Dashboard simulator', 'Handbook'],
        'videos': ['Introduction to Driver Responsibilities', 'Understanding Your Vehicle'],
        'activities': ['Group discussion on road user experiences', 'Dashboard warning light identification quiz'],
        'assessment': 'Verbal Q&A on cockpit drill and dashboard lights',
    },
    {
        'sat': 'Saturday 2 \nRoad Signs & Traffic Rules',
        'theme': 'Road Signs, Markings & Highway Code',
        'topics': [
            'Road sign categories (regulatory, warning, informative)',
            'Road markings (centre lines, edge lines, hatched areas)',
            'Traffic light sequences and meanings',
            'Speed limits for different vehicle types and roads',
            'Hand signals for drivers and cyclists',
            'Highway Code rules and best practices',
        ],
        'materials': ['Highway Code booklets', 'Road sign flash cards', 'Interactive sign quiz platform'],
        'videos': ['Road Signs Explained', 'Understanding Traffic Lights & Signals'],
        'activities': ['Road sign matching game', 'Speed limit calculation exercise'],
        'assessment': 'Written road sign identification quiz',
    },
    {
        'sat': 'Saturday 3 \nParking, Junctions & Hazard Perception',
        'theme': 'Parking Theory, Junctions & Defensive Driving',
        'topics': [
            'Parallel parking theory and reference points',
            'Roundabout rules and lane selection',
            'Junction types and right-of-way rules',
            'Mirror use and distance judgement',
            'Blind spot awareness and checks',
            'Defensive driving principles and hazard perception',
        ],
        'materials': ['Hazard perception video clips', 'Junction diagram handout', 'Roundabout simulator'],
        'videos': ['Mastering Junctions', 'Defensive Driving & Hazard Perception'],
        'activities': ['Hazard perception video analysis', 'Junction priority exercise'],
        'assessment': 'Hazard perception spot-the-hazard exercise',
    },
    {
        'sat': 'Saturday 4 \nEmergencies & Examination Preparation',
        'theme': 'Emergency Procedures & Mock Theory Examination',
        'topics': [
            'Brake failure procedures and safe stopping',
            'Tyre burst handling and control',
            'Accident procedures (stop, assess, document, report)',
            'First aid basics for drivers',
            'Comprehensive theory revision',
            'Mock theory examination',
        ],
        'materials': ['First aid kit demonstration', 'Emergency procedure cards', 'Theory test papers'],
        'videos': ['Handling Emergencies', 'First Aid for Drivers'],
        'activities': ['Mock theory examination', 'Emergency scenario role-play'],
        'assessment': 'Mock theory examination (pass mark: 43/50)',
    },
]

for sat in saturdays:
    doc.add_heading(sat['sat'], level=2)

    p_theme = doc.add_paragraph()
    run_theme = p_theme.add_run(f'Theme: {sat["theme"]}')
    run_theme.bold = True
    run_theme.font.size = Pt(11)
    run_theme.font.name = 'Calibri'
    run_theme.font.color.rgb = TEAL

    # Topics
    p_t = doc.add_paragraph()
    run_t = p_t.add_run('Topics Covered:')
    run_t.bold = True
    run_t.font.size = Pt(11)
    run_t.font.name = 'Calibri'
    for topic in sat['topics']:
        p_topic = doc.add_paragraph()
        run_b = p_topic.add_run('\u2022  ')
        run_b.font.name = 'Calibri'
        run_t2 = p_topic.add_run(topic)
        run_t2.font.name = 'Calibri'
        p_topic.paragraph_format.space_after = Pt(1)
        p_topic.paragraph_format.space_before = Pt(0)

    # Materials, Videos, Activities, Assessment table
    sub_headers = ['Category', 'Details']
    sub_rows = [
        ['Teaching Materials', '\n'.join(f'\u2022 {m}' for m in sat['materials'])],
        ['Videos', '\n'.join(f'\u2022 {v}' for v in sat['videos'])],
        ['Activities', '\n'.join(f'\u2022 {a}' for a in sat['activities'])],
        ['Assessment', sat['assessment']],
    ]
    add_styled_table(doc, sub_headers, sub_rows, col_widths=[4, 14])
    doc.add_paragraph()

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# SECTION 6: DRIVING COMPETENCIES MATRIX
# ═══════════════════════════════════════════════════════════════════

doc.add_heading('6. Driving Competencies Matrix', level=1)

p_comp_intro = doc.add_paragraph()
run_ci = p_comp_intro.add_run(
    'The following matrix maps every driving competency across the 20-day curriculum, indicating when each '
    'skill is Introduced (I), Practised (P), Assessed (A), and Mastered (M). This matrix ensures complete '
    'coverage and identifies any gaps in the curriculum.'
)
run_ci.font.size = Pt(11)
run_ci.font.name = 'Calibri'

doc.add_paragraph()

competencies_matrix = [
    ['Cockpit Drill (DSSSM)', '1', '2\u20135', '5', '5'],
    ['Vehicle Inspection', '2', '3\u20135', '5', '5'],
    ['Vehicle Control', '3', '4\u20136', '5', '10'],
    ['Clutch Control (Manual)', '3', '4\u20138', '10', '15'],
    ['Gear Changing (Manual)', '3', '4\u201312', '10', '15'],
    ['Automatic Controls (Auto)', '3', '4\u20138', '10', '15'],
    ['Observation (MSM Routine)', '4', '5\u201312', '12', '15'],
    ['MSM Routine', '6', '7\u201312', '12', '15'],
    ['Road Positioning', '6', '7\u201311', '11', '16'],
    ['Junctions', '7', '8\u201312', '12', '16'],
    ['Roundabouts', '12', '13\u201317', '17', '19'],
    ['Reverse Driving', '9', '10\u201314', '14', '18'],
    ['Reverse Parking', '9', '10\u201313', '13', '18'],
    ['Bay Parking', '10', '11\u201314', '14', '18'],
    ['Angle Parking', '10', '11\u201314', '14', '18'],
    ['Parallel Parking', '13', '14\u201317', '18', '18'],
    ['Hill Starts', '8', '9\u201314', '14', '16'],
    ['Hill Parking', '8', '9\u201314', '14', '16'],
    ['Turning Around (3-Point/U-Turn)', '14', '15\u201317', '17', '19'],
    ['Mirror Distance Judgement', '6', '7\u201316', '16', '19'],
    ['Safe Overtaking', '14', '15\u201317', '17', '19'],
    ['Defensive Driving', '16', '17\u201319', '19', '20'],
    ['Hazard Perception', '15', '16\u201319', '19', '20'],
    ['Emergency Stop', '15', '16\u201319', '19', '20'],
    ['Night Driving Awareness', '16', '17\u201319', '19', '20'],
    ['Rain/Wet Weather Driving', '16', '17\u201319', '19', '20'],
    ['Independent Driving', '17', '18\u201319', '19', '20'],
]

matrix_headers = ['Competency', 'I', 'P', 'A', 'M']
add_styled_table(doc, matrix_headers, competencies_matrix, col_widths=[6, 2, 2, 2, 2])

doc.add_paragraph()

# Legend
p_leg = doc.add_paragraph()
run_leg = p_leg.add_run('Legend:  I = Introduced  |  P = Practised  |  A = Assessed  |  M = Mastered')
run_leg.bold = True
run_leg.font.size = Pt(10)
run_leg.font.color.rgb = TEAL
run_leg.font.name = 'Calibri'

add_icon_text(doc, '\u2714\ufe0f', ' All 27 competencies are covered across the 20-day curriculum.')
add_icon_text(doc, '\u2714\ufe0f', ' Competencies are revisited at least twice before formal assessment.')
add_icon_text(doc, '\u2714\ufe0f', ' Mastery (M) is achieved only after successful formal assessment.')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# SECTION 7: LESSON VIDEO LIBRARY
# ═══════════════════════════════════════════════════════════════════

doc.add_heading('7. Lesson Video Library', level=1)

p_vid_intro = doc.add_paragraph()
run_vi = p_vid_intro.add_run(
    'Each practical lesson is supported by a short instructional video (2\u20133 minutes) that students watch '
    'at the beginning of the session. Videos are accessed via tablet or in-vehicle display. '
    'QR codes are available for students to scan and review lessons at home.'
)
run_vi.font.size = Pt(11)
run_vi.font.name = 'Calibri'

doc.add_paragraph()

video_library = [
    ['1', 'Vehicle Orientation & Controls', '2:30', '[QR Placeholder]', 'Identify all vehicle controls and dashboard indicators'],
    ['2', 'Daily Vehicle Inspection', '2:45', '[QR Placeholder]', 'Perform complete daily vehicle safety inspection'],
    ['3', 'Clutch Control Basics / Automatic Controls', '3:00', '[QR Placeholder]', 'Develop smooth clutch or pedal control'],
    ['4', 'Steering Techniques', '2:30', '[QR Placeholder]', 'Apply correct pull-push steering method'],
    ['5', 'Week 1 Consolidation', '3:00', '[QR Placeholder]', 'Reinforce all Week 1 skills and competencies'],
    ['6', 'Moving Off & Stopping', '2:45', '[QR Placeholder]', 'Execute safe moving off and stopping procedures'],
    ['7', 'Approaching Junctions', '3:00', '[QR Placeholder]', 'Navigate junctions with correct observation and gear'],
    ['8', 'Hill Starts & Incline Parking', '3:00', '[QR Placeholder]', 'Perform hill starts and parking on inclines'],
    ['9', 'Reverse Driving & Parking', '3:00', '[QR Placeholder]', 'Reverse confidently and park between cars'],
    ['10', 'Bay & Angle Parking', '2:45', '[QR Placeholder]', 'Master forward/reverse bay and angle parking'],
    ['11', 'Road Positioning & Lane Discipline', '2:30', '[QR Placeholder]', 'Maintain correct road position at all times'],
    ['12', 'Navigating Roundabouts', '3:00', '[QR Placeholder]', 'Navigate single and multi-lane roundabouts'],
    ['13', 'Parallel Parking Step by Step', '3:00', '[QR Placeholder]', 'Execute parallel parking using reference points'],
    ['14', 'Turning Around & Overtaking', '2:45', '[QR Placeholder]', 'Three-point turn, U-turn, and safe overtaking'],
    ['15', 'Emergency Procedures', '3:00', '[QR Placeholder]', 'Respond to emergencies: stop, avoid, and recover'],
    ['16', 'Defensive Driving Techniques', '2:45', '[QR Placeholder]', 'Apply defensive driving and hazard perception'],
    ['17', 'Mixed Road Driving Strategies', '3:00', '[QR Placeholder]', 'Drive confidently on highways, urban, and rural roads'],
    ['18', 'Parallel Parking \u2013 Assessment Prep', '2:30', '[QR Placeholder]', 'Prepare for parallel parking assessment'],
    ['19', 'Mock Test Preparation', '3:00', '[QR Placeholder]', 'Understand test procedures and manage stress'],
    ['20', 'Final Assessment Review', '3:00', '[QR Placeholder]', 'Comprehensive review of all driving competencies'],
]

video_headers = ['Lesson', 'Video Title', 'Duration', 'QR Code', 'Learning Outcome']
add_styled_table(doc, video_headers, video_library, col_widths=[1.5, 5, 1.5, 3, 7])

add_instructor_tip(doc, 'Ensure the video player is ready before the lesson begins. The 2\u20133 minute video should not cut into the 20-minute practical session time.')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# SECTION 8: INSTRUCTOR DAILY ASSESSMENT FORM
# ═══════════════════════════════════════════════════════════════════

doc.add_heading('8. Instructor Daily Assessment Form', level=1)

p_form_intro = doc.add_paragraph()
run_fi = p_form_intro.add_run(
    'This form is completed by the instructor at the end of every practical lesson. It serves as the primary '
    'record of student progress and feeds directly into the CRM system. Forms should be completed digitally '
    'in the CRM or printed for physical records.'
)
run_fi.font.size = Pt(11)
run_fi.font.name = 'Calibri'

doc.add_paragraph()

# Assessment form as styled table
form_headers = ['Field', 'Entry']
form_rows = [
    ['Student Name', '__________________________________'],
    ['Instructor', '__________________________________'],
    ['Vehicle Registration', '__________________________________'],
    ['Transmission', '\u2610  Manual    \u2610  Automatic'],
    ['Lesson Number', '______  of 20'],
    ['Date', '__________________'],
    ['Start Time', '__________________'],
    ['End Time', '__________________'],
    ['Distance Driven', '__________ km'],
    ['Video Watched', '\u2610  Yes    \u2610  No'],
    ['Vehicle Inspection Completed', '\u2610  Yes    \u2610  No'],
    ['Cockpit Drill Completed', '\u2610  Yes    \u2610  No'],
]
add_styled_table(doc, form_headers, form_rows, col_widths=[5, 13])

doc.add_paragraph()

p_comp_title = doc.add_paragraph()
run_ct = p_comp_title.add_run('Competencies Achieved (check all that apply):')
run_ct.bold = True
run_ct.font.size = Pt(11)
run_ct.font.name = 'Calibri'

comp_items = [
    'Cockpit Drill (DSSSM)', 'Vehicle Inspection', 'Vehicle Control',
    'Clutch Control (Manual)', 'Gear Changing', 'Automatic Controls',
    'Observation (MSM)', 'Road Positioning', 'Junction Approach',
    'Roundabout Navigation', 'Reverse Driving', 'Reverse Parking',
    'Bay Parking', 'Angle Parking', 'Parallel Parking',
    'Hill Start', 'Hill Parking', 'Turning Around',
    'Safe Overtaking', 'Defensive Driving', 'Hazard Perception',
    'Emergency Stop', 'Independent Driving',
]

for i in range(0, len(comp_items), 2):
    p_row = doc.add_paragraph()
    run_c1 = p_row.add_run(f'\u2610  {comp_items[i]}')
    run_c1.font.size = Pt(10)
    run_c1.font.name = 'Calibri'
    run_c1.font.color.rgb = BLACK
    run_sp = p_row.add_run('    \t')
    if i + 1 < len(comp_items):
        run_c2 = p_row.add_run(f'\u2610  {comp_items[i + 1]}')
        run_c2.font.size = Pt(10)
        run_c2.font.name = 'Calibri'
        run_c2.font.color.rgb = BLACK

doc.add_paragraph()

asses_headers = ['Section', 'Notes']
asses_rows = [
    ['Strengths', '__________________________________\n__________________________________\n__________________________________'],
    ['Weaknesses / Areas for Improvement', '__________________________________\n__________________________________\n__________________________________'],
    ['Instructor Recommendations', '__________________________________\n__________________________________\n__________________________________'],
]
add_styled_table(doc, asses_headers, asses_rows, col_widths=[5, 13])

doc.add_paragraph()

# Signatures
sig_headers = ['Role', 'Name', 'Signature', 'Date']
sig_rows = [
    ['Instructor', '__________________', '__________________', '____________'],
    ['Student', '__________________', '__________________', '____________'],
]
add_styled_table(doc, sig_headers, sig_rows, col_widths=[3, 5, 5, 3])

doc.add_paragraph()

add_instructor_tip(doc, 'Complete this form immediately after each lesson while observations are fresh. Be specific in the strengths and weaknesses sections \u2014 general feedback is not as helpful for student development.')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# SECTION 9: CRM MAPPING
# ═══════════════════════════════════════════════════════════════════

doc.add_heading('9. CRM Functional Mapping', level=1)

p_crm_intro = doc.add_paragraph()
run_crm = p_crm_intro.add_run(
    'This section maps every element of the Fance Driving School curriculum to the corresponding fields and '
    'features in the Driving School CRM. This serves as the functional specification for CRM implementation '
    'and ensures that the software supports all curriculum requirements.'
)
run_crm.font.size = Pt(11)
run_crm.font.name = 'Calibri'

doc.add_paragraph()

crm_mapping = [
    ['Package Purchased', 'CRM Products / Packages module', 'Package selected defines transmission type, training flags, and duration fields'],
    ['Training Template Selected', 'CRM Lesson Plan Templates', 'Template selected per transmission type; populates 20-day lesson structure'],
    ['Student Training Plan Generated', 'CRM Client Lesson Plan', 'Auto-generates 20 lessons from template, assigned to student'],
    ['Lesson Number', 'CRM Lesson.number (integer)', '1\u201320, inherited from lesson plan template'],
    ['Lesson Title', 'CRM Lesson.title (text)', 'Maps to curriculum lesson name'],
    ['Competencies', 'CRM Lesson.skills_achieved (JSONB)', 'Skill name, competency level 1\u20135, achieved boolean'],
    ['Checklist', 'CRM Lesson.instructor_checklist (JSONB)', 'Predefined checklist items per lesson template'],
    ['Video', 'CRM Lesson.video_url + video_cached', 'URL to lesson video; cached flag for offline viewing'],
    ['Video Watched', 'CRM TrainingSession video fields', 'Tracks video viewed status, cache/invalidate actions'],
    ['Assessment', 'CRM Lesson.status (enum)', 'completed / partial / not_started / skipped'],
    ['Status', 'CRM Lesson.status + progress tracking', 'Derived from checklist completion percentage'],
    ['Session Duration', 'CRM TrainingSession.timer_seconds', 'Timer tracks actual driving time (max 30 min)'],
    ['Distance', 'CRM Lesson odometer reading', 'Track distance driven during session (max 3 km)'],
    ['Start / End Time', 'CRM TrainingSession.started_at + timer', 'Start time, accumulated timer, end time'],
    ['Instructor Notes', 'CRM Lesson.notes (text)', 'Free-text instructor observations per lesson'],
    ['Strengths', 'CRM Lesson.strengths (text)', 'Key strengths identified during assessment'],
    ['Weaknesses', 'CRM Lesson.weaknesses (text)', 'Areas requiring improvement'],
    ['Recommendations', 'CRM Lesson.recommendations (text)', 'Instructor recommendations for next lesson'],
    ['Instructor Signature', 'CRM User authentication', 'Instructor confirmed via authenticated session'],
    ['Student Signature', 'CRM digital signature field', 'Student acknowledgement captured digitally'],
    ['Repeat Required', 'CRM Lesson.repeat_required (boolean)', 'Flag for lesson to be repeated before progression'],
    ['Next Lesson', 'CRM LessonPlan.next_lesson_id (FK)', 'Auto-computed based on completion status'],
    ['Progress Dashboard', 'CRM Consultation profile page', 'Real-time progress bar per cart item / student'],
]

crm_headers = ['Curriculum Element', 'CRM Field / Module', 'Implementation Notes']
add_styled_table(doc, crm_headers, crm_mapping, col_widths=[4, 5, 9])

add_icon_text(doc, '\u2714\ufe0f', ' All 22 curriculum elements are mapped to corresponding CRM fields.')
add_icon_text(doc, '\u2714\ufe0f', ' The CRM is designed to mirror the curriculum workflow exactly.')
add_icon_text(doc, '\u2714\ufe0f', ' Training sessions, skills, and lesson plan modules are already implemented in the current CRM.')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# SECTION 10: CRM WORKFLOW
# ═══════════════════════════════════════════════════════════════════

doc.add_heading('10. CRM Workflow Diagram', level=1)

p_workflow_intro = doc.add_paragraph()
run_wf = p_workflow_intro.add_run(
    'The following workflow illustrates the complete student journey from package purchase through to '
    'licensing test readiness, as managed by the Driving School CRM. Each step corresponds to specific '
    'CRM modules and actions.'
)
run_wf.font.size = Pt(11)
run_wf.font.name = 'Calibri'

doc.add_paragraph()

workflow_steps = [
    'PACKAGE PURCHASED',
    'TRAINING TEMPLATE SELECTED',
    'STUDENT TRAINING PLAN GENERATED',
    'DAILY LESSON',
    'VIDEO VIEWED',
    'INSTRUCTOR CHECKLIST',
    'COMPETENCIES ACHIEVED',
    'LESSON COMPLETED',
    'PROGRESS DASHBOARD UPDATED',
    'STUDENT READY FOR LICENSING TEST',
]

for i, step in enumerate(workflow_steps):
    p_step = doc.add_paragraph()
    p_step.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_step.paragraph_format.space_before = Pt(2)
    p_step.paragraph_format.space_after = Pt(2)

    box_table = doc.add_table(rows=1, cols=1)
    box_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = box_table.rows[0].cells[0]
    set_cell_shading(cell, '008080')
    set_cell_text(cell, step, bold=True, color=WHITE, size=Pt(12), alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_vertical_alignment(cell, "center")
    for row in box_table.rows:
        for c in row.cells:
            c.width = Cm(18)

    if i < len(workflow_steps) - 1:
        p_arrow = doc.add_paragraph()
        p_arrow.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_arrow.paragraph_format.space_before = Pt(4)
        p_arrow.paragraph_format.space_after = Pt(4)
        run_arrow = p_arrow.add_run('\u2193')
        run_arrow.font.size = Pt(20)
        run_arrow.font.color.rgb = GOLD
        run_arrow.font.name = 'Segoe UI Symbol'

doc.add_paragraph()

add_instructor_tip(doc, 'The CRM workflow ensures no step is skipped. If a lesson is not completed (competencies not achieved), the workflow loops back to "Daily Lesson" until the student passes.')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# SECTION 11: APPENDIX
# ═══════════════════════════════════════════════════════════════════

doc.add_heading('11. Appendix', level=1)

# ── Appendix A: Daily Vehicle Inspection Checklist ──
doc.add_heading('Appendix A: Daily Vehicle Inspection Checklist', level=2)

p_app_a = doc.add_paragraph()
run_aa = p_app_a.add_run(
    'This checklist must be completed by the student under instructor supervision at the start of each practical lesson.'
)
run_aa.font.size = Pt(11)
run_aa.font.name = 'Calibri'
run_aa.italic = True

doc.add_paragraph()

inspection_items = [
    ['Exterior Walk-Around', ''],
    ['Tyre condition (cracks, bulges, tread depth)', '\u2610'],
    ['Tyre pressure (visual inspection)', '\u2610'],
    ['Headlights (low and high beam)', '\u2610'],
    ['Tail lights', '\u2610'],
    ['Brake lights', '\u2610'],
    ['Left turn indicator', '\u2610'],
    ['Right turn indicator', '\u2610'],
    ['Hazard warning lights', '\u2610'],
    ['Number plate (clean and legible)', '\u2610'],
    ['Engine Compartment', ''],
    ['Engine oil level (dipstick check)', '\u2610'],
    ['Coolant level (reservoir check)', '\u2610'],
    ['Brake fluid level', '\u2610'],
    ['Windshield washer fluid level', '\u2610'],
    ['Interior Checks', ''],
    ['Horn operation', '\u2610'],
    ['Windshield wipers (front and rear)', '\u2610'],
    ['Washer jets (front and rear)', '\u2610'],
    ['Seat adjustment (all positions)', '\u2610'],
    ['Steering wheel adjustment', '\u2610'],
    ['Mirror adjustment (interior and side)', '\u2610'],
    ['Seatbelt (driver and passenger)', '\u2610'],
    ['Handbrake operation', '\u2610'],
    ['Gear selection (all positions)', '\u2610'],
    ['Dashboard warning lights (self-check on ignition)', '\u2610'],
]

insp_headers = ['Inspection Item', 'Check']
add_styled_table(doc, insp_headers, inspection_items, col_widths=[16, 2])

doc.add_paragraph()
p_sig_a = doc.add_paragraph()
run_sa = p_sig_a.add_run('Student Signature: ________________________  Date: ____________')
run_sa.font.size = Pt(10)
run_sa.font.name = 'Calibri'

p_sig_a2 = doc.add_paragraph()
run_sa2 = p_sig_a2.add_run('Instructor Signature: ______________________  Date: ____________')
run_sa2.font.size = Pt(10)
run_sa2.font.name = 'Calibri'

doc.add_paragraph()

add_safety_note(doc, 'Never operate a vehicle that fails any safety inspection item. Any failed item must be reported to the fleet manager immediately and rectified before the vehicle is used for training.')

doc.add_page_break()

# ── Appendix B: Cockpit Drill Checklist ──
doc.add_heading('Appendix B: Cockpit Drill Checklist (DSSSM)', level=2)

p_app_b = doc.add_paragraph()
run_ab = p_app_b.add_run(
    'The DSSSM (Doors, Seat, Steering, Seatbelt, Mirrors) cockpit drill must be performed at the start of every practical lesson.'
)
run_ab.font.size = Pt(11)
run_ab.font.name = 'Calibri'
run_ab.italic = True

doc.add_paragraph()

cockpit_items = [
    ['D \u2014 DOORS', 'Ensuring all doors are properly closed and locked; red lights on dashboard off'],
    ['\u2714  All doors closed securely', '\u2610'],
    ['\u2714  Dashboard door indicator off', '\u2610'],
    ['S \u2014 SEAT', 'Adjusting seat for reach to pedals, steering wheel visibility, and comfort'],
    ['\u2714  Seat distance: pedals fully depressed without stretching', '\u2610'],
    ['\u2714  Seat height: clear view of road and dashboard', '\u2610'],
    ['\u2714  Seat backrest: slight recline, arms slightly bent on steering wheel', '\u2610'],
    ['S \u2014 STEERING', 'Adjusting steering wheel for comfort and control'],
    ['\u2714  Steering wheel adjusted for reach', '\u2610'],
    ['\u2714  Hands at 10-to-2 or 9-to-3 position', '\u2610'],
    ['\u2714  Full steering rotation without stretching', '\u2610'],
    ['S \u2014 SEATBELT', 'Fastening seatbelt correctly'],
    ['\u2714  Seatbelt fastened and snug', '\u2610'],
    ['\u2714  No twists in belt', '\u2610'],
    ['\u2714  Belt across shoulder and lap correctly', '\u2610'],
    ['M \u2014 MIRRORS', 'Adjusting all mirrors for maximum visibility'],
    ['\u2714  Rear-view mirror: full rear window view', '\u2610'],
    ['\u2714  Left side mirror: minimal car body, maximum road view', '\u2610'],
    ['\u2714  Right side mirror: minimal car body, maximum road view', '\u2610'],
]

cockpit_headers = ['Step', 'Check']
add_styled_table(doc, cockpit_headers, cockpit_items, col_widths=[10, 8])

doc.add_paragraph()

add_instructor_tip(doc, 'Teach the DSSSM mnemonic from the very first lesson. Students should be able to recite and perform the drill without prompting by Day 5.')

doc.add_page_break()

# ── Appendix C: Parking Reference Points ──
doc.add_heading('Appendix C: Parking Reference Points', level=2)

p_app_c = doc.add_paragraph()
run_ac = p_app_c.add_run(
    'The following reference points are specific to the training vehicle and must be taught consistently across all instructors. '
    'Reference points may vary slightly between Manual and Automatic transmission models due to different vehicle dimensions.'
)
run_ac.font.size = Pt(11)
run_ac.font.name = 'Calibri'
run_ac.italic = True

doc.add_paragraph()

ref_headers = ['Parking Manoeuvre', 'Reference Point', 'Description']
ref_rows = [
    ['Forward Bay Parking', 'Shoulder alignment', 'Stop when your shoulder aligns with the second line of the bay'],
    ['Forward Bay Parking', 'Full lock point', 'Turn full lock when mirror passes the bay line'],
    ['Forward Bay Parking', 'Centre reference', 'Centre bonnet emblem aligns with bay centre line'],
    ['Reverse Bay Parking', 'Mirror alignment', 'Rear door handle aligns with adjacent car bumper'],
    ['Reverse Bay Parking', 'Full lock point', 'Turn full lock when front wheel aligns with bay line'],
    ['Reverse Bay Parking', 'Straighten point', 'Straighten wheels when car is parallel to lines'],
    ['Angle Parking', 'Bonnet corner', 'Bonnet corner aligns with parking bay line'],
    ['Angle Parking', 'Full lock point', 'Turn full lock when front wheels reach bay entry'],
    ['Parallel Parking', 'Bumper alignment', 'Align rear bumper with rear bumper of front car'],
    ['Parallel Parking', 'Full lock left', 'Turn full lock left when rear wheel aligns with front car bumper'],
    ['Parallel Parking', 'Straighten point', 'Straighten wheels when kerb appears at base of rear window'],
    ['Parallel Parking', 'Full lock right', 'Turn full lock right when front of car clears rear car'],
    ['Parallel Parking', 'Final position', 'Stop when car is parallel, 20\u201330 cm from kerb'],
]

add_styled_table(doc, ref_headers, ref_rows, col_widths=[4.5, 4, 9.5])

add_instructor_tip(doc, 'Mark reference points with small stickers or tape on the vehicle during the first week. Remove them by Week 3 to encourage spatial awareness.')

doc.add_page_break()

# ── Appendix D: Manual vs Automatic Comparison ──
doc.add_heading('Appendix D: Manual vs Automatic Transmission Comparison', level=2)

p_app_d = doc.add_paragraph()
run_ad = p_app_d.add_run(
    'The following table compares the key differences between Manual and Automatic transmission training. '
    'Students should choose their preferred transmission type before enrolment.'
)
run_ad.font.size = Pt(11)
run_ad.font.name = 'Calibri'
run_ad.italic = True

doc.add_paragraph()

compare_headers = ['Aspect', 'Manual Transmission', 'Automatic Transmission']
compare_rows = [
    ['Pedals', '3 pedals: Clutch, Brake, Accelerator', '2 pedals: Brake, Accelerator'],
    ['Gear Selection', '6-speed + Reverse (H-pattern)', 'PRNDL (Park, Reverse, Neutral, Drive, Low)'],
    ['Clutch Control', 'Required for all gear changes and moving off', 'Not required'],
    ['Hill Start', 'Handbrake + clutch bite point method', 'Handbrake + accelerator (simpler)'],
    ['Learning Curve', 'Steeper \u2014 requires coordination', 'Gentler \u2014 focus on steering and observation'],
    ['Fuel Efficiency', 'Typically higher with skilled driver', 'Modern automatics are comparable'],
    ['Driving License', 'Qualifies for Manual license', 'Limited to Automatic license'],
    ['Vehicle Availability', 'Broader selection including older cars', 'Increasingly common, especially EV/hybrid'],
    ['Training Duration', 'May require additional lessons', 'Generally fewer lessons needed'],
    ['Lesson Cost', 'Same as Automatic', 'Same as Manual'],
]

add_styled_table(doc, compare_headers, compare_rows, col_widths=[3.5, 7, 7.5])

doc.add_paragraph()

add_icon_text(doc, '\u2714\ufe0f', ' Students licensed on Manual transmission can drive both Manual and Automatic vehicles.')
add_icon_text(doc, '\u2714\ufe0f', ' Students licensed on Automatic transmission can only drive Automatic vehicles.')

doc.add_page_break()

# ── Appendix E: Road Test Readiness Checklist ──
doc.add_heading('Appendix E: Road Test Readiness Checklist', level=2)

p_app_e = doc.add_paragraph()
run_ae = p_app_e.add_run(
    'This checklist is used by the instructor to determine whether a student is ready to sit the official practical driving test. '
    'All items must be checked before scheduling the test.'
)
run_ae.font.size = Pt(11)
run_ae.font.name = 'Calibri'
run_ae.italic = True

doc.add_paragraph()

roadtest_items = [
    ['Complete 20-Day Curriculum', 'All 20 lessons completed in the CRM', '\u2610'],
    ['Parallel Parking Assessment', 'Passed Day 18 assessment (Training Pitch)', '\u2610'],
    ['Mock Test Pass', 'Passed Day 19 Mock Practical Test', '\u2610'],
    ['Final Assessment Pass', 'Passed Day 20 Final Assessment', '\u2610'],
    ['Theory Examination Pass', 'Passed Mock Theory Exam (43/50 minimum)', '\u2610'],
    ['Independent Driving', 'Drives independently without instructor prompting', '\u2610'],
    ['Hazard Perception', 'Demonstrates consistent hazard awareness', '\u2610'],
    ['Defensive Driving', 'Applies defensive driving habits consistently', '\u2610'],
    ['Manoeuvres Competence', 'All manoeuvres performed correctly on demand', '\u2610'],
    ['Emergency Stop', 'Controlled emergency stop demonstrated', '\u2610'],
    ['Traffic Rules Knowledge', 'Demonstrates full knowledge of highway code', '\u2610'],
    ['Test Nerves Management', 'Student reports confidence and readiness', '\u2610'],
    ['Vehicle Familiarisation', 'Comfortable with test vehicle controls', '\u2610'],
    ['Documents Ready', 'Valid learner permit, licence, identification', '\u2610'],
]

roadtest_headers = ['Readiness Criteria', 'Details', 'Check']
add_styled_table(doc, roadtest_headers, roadtest_items, col_widths=[5, 9, 2])

doc.add_paragraph()

p_rt_sig = doc.add_paragraph()
run_rts = p_rt_sig.add_run(
    'I confirm that the above-named student has met all readiness criteria and is recommended for the official road test.'
)
run_rts.font.size = Pt(11)
run_rts.font.name = 'Calibri'

p_rt_sig2 = doc.add_paragraph()
run_rts2 = p_rt_sig2.add_run('Instructor Name: ____________________  Signature: ____________________  Date: ____________')
run_rts2.font.size = Pt(10)
run_rts2.font.name = 'Calibri'

doc.add_page_break()

# ── Appendix F: Instructor Weekly Report ──
doc.add_heading('Appendix F: Instructor Weekly Report', level=2)

p_app_f = doc.add_paragraph()
run_af = p_app_f.add_run(
    'Completed by the instructor at the end of each training week and submitted to the Training Manager.'
)
run_af.font.size = Pt(11)
run_af.font.name = 'Calibri'
run_af.italic = True

doc.add_paragraph()

weekly_headers = ['Field', 'Details']
weekly_rows = [
    ['Instructor Name', '__________________________'],
    ['Week Number', 'Week ____  of 4'],
    ['Reporting Period', '__________________  to  __________________'],
    ['Students Taught This Week', '__________'],
    ['Lessons Delivered', '__________ / 25 (5 students \u00d7 5 days)'],
    ['Lessons Completed (Pass)', '__________'],
    ['Lessons Repeated', '__________'],
    ['Lessons Outstanding', '__________'],
    ['Students Ready for Next Week', '__________'],
    ['Vehicles Used', '__________________________'],
    ['Vehicle Issues Reported', '__________________________'],
]
add_styled_table(doc, weekly_headers, weekly_rows, col_widths=[6, 12])

doc.add_paragraph()

p_wn = doc.add_paragraph()
run_wn = p_wn.add_run('Weekly Narrative Report:')
run_wn.bold = True
run_wn.font.size = Pt(11)
run_wn.font.name = 'Calibri'

narrative_fields = [
    'Key achievements this week:',
    'Challenges encountered:',
    'Student progress notes:',
    'Vehicle / equipment issues:',
    'Recommendations for next week:',
    'Training Manager comments:',
]

for field in narrative_fields:
    p_nf = doc.add_paragraph()
    run_nf = p_nf.add_run(f'{field}\n{"_" * 80}')
    run_nf.font.size = Pt(10)
    run_nf.font.name = 'Calibri'

doc.add_page_break()

# ── Appendix G: Student Progress Tracker ──
doc.add_heading('Appendix G: Student Progress Tracker', level=2)

p_app_g = doc.add_paragraph()
run_ag = p_app_g.add_run(
    'This tracker summarises a single student\'s progress across the entire 20-day programme. '
    'It is updated after every lesson and forms the basis of the CRM progress dashboard.'
)
run_ag.font.size = Pt(11)
run_ag.font.name = 'Calibri'
run_ag.italic = True

doc.add_paragraph()

student_info_headers = ['Field', 'Details']
student_info_rows = [
    ['Student Name', '__________________________'],
    ['Transmission', '\u2610  Manual    \u2610  Automatic'],
    ['Start Date', '__________________________'],
    ['Target Completion', '__________________________'],
]
add_styled_table(doc, student_info_headers, student_info_rows, col_widths=[5, 13])

doc.add_paragraph()

lesson_names = [
    'Vehicle Orientation', 'Vehicle Checks', 'Clutch Control / Auto Controls', 'Steering', 'Consolidation',
    'Moving Off', 'Junctions', 'Hill Starts & Parking', 'Reverse Driving & Parking', 'Bay & Angle Parking',
    'Road Positioning', 'Roundabouts', 'Parallel Parking (Pitch)', 'Turning Around & Overtaking', 'Emergency Procedures',
    'Defensive Driving', 'Mixed Road Driving', 'Parallel Parking Assessment', 'Mock Test', 'Final Assessment',
]

progress_headers_short = ['Day', 'Lesson', 'Date', 'Status', 'Competencies', 'Min', 'km', 'Video', 'Repeat', 'Instructor Notes']
progress_rows = [
    [str(i), lesson_names[i-1], '', '\u2610', '\u2610', '', '', '\u2610', '\u2610', '']
    for i in range(1, 21)
]
add_styled_table(doc, progress_headers_short, progress_rows, col_widths=[1, 4, 2, 1.5, 1.5, 1, 1, 1, 1, 3])

doc.add_paragraph()

# Summary
p_summary = doc.add_paragraph()
run_sum = p_summary.add_run('PROGRAMME SUMMARY')
run_sum.bold = True
run_sum.font.size = Pt(12)
run_sum.font.color.rgb = TEAL
run_sum.font.name = 'Calibri'

summary_headers = ['Metric', 'Value']
summary_rows = [
    ['Total Lessons', '20 of 20'],
    ['Lessons Passed', '______'],
    ['Lessons Repeated', '______'],
    ['Total Driving Time', '______ minutes'],
    ['Total Distance', '______ km'],
    ['Theory Sessions Attended', '______ of 4'],
    ['Overall Status', '\u2610  In Progress    \u2610  Completed    \u2610  Graduated'],
]
add_styled_table(doc, summary_headers, summary_rows, col_widths=[6, 12])

doc.add_paragraph()

p_final = doc.add_paragraph()
p_final.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_final.paragraph_format.space_before = Pt(20)
run_final = p_final.add_run('\u2014  End of Document  \u2014')
run_final.bold = True
run_final.font.size = Pt(14)
run_final.font.color.rgb = GOLD
run_final.font.name = 'Calibri'

# ─── Save the document ──────────────────────────────────────────
output_path = '/Users/mac/Documents/Projects/ai/driving_school_crm_2/Fance_Driving_School_Curriculum_v1.0.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')
